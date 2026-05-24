import json
import logging
import os
from collections.abc import AsyncIterator
from typing import Literal

import httpx
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.middleware.auth import require_service_auth
from app.models.schemas import (
    BrdSectionScore,
    BrdTemplateScore,
    ChatRequest,
    ChatResponse,
    CvParseRequest,
    CvParseResponse,
    CvParsedData,
    GenerateBrdRequest,
    GenerateBrdResponse,
    GeneratePrdRequest,
    GeneratePrdResponse,
    MatchingRequest,
    MatchingResponse,
    ParseSpecData,
    ParseSpecRequest,
    ParseSpecResponse,
)
from app.services.nats_client import publish_event

logger = logging.getLogger(__name__)

router = APIRouter()

TENSORZERO_URL = os.getenv("TENSORZERO_API_URL", "http://localhost:3333")
PROJECT_SERVICE_URL = os.getenv("PROJECT_SERVICE_URL", "http://localhost:3002")


def _service_auth_secret() -> str:
    """Outgoing X-Service-Auth secret. Read at call time so tests can override."""
    return os.getenv("SERVICE_AUTH_SECRET", "")


def calculate_completeness(messages: list) -> int:
    """Score chat conversation against BRD template info requirements (sections B-N).

    Each check maps to a BRD template section that needs real data from the client.
    Score = covered_checks / total_checks * 100.
    """
    user_messages = [m.content.lower() for m in messages if m.role == "user"]
    if not user_messages:
        return 0

    all_text = " ".join(user_messages)

    # Section B — Executive Summary: project description present
    has_description = len(all_text) > 80

    # Section C — Problem Statement: pain points or motivation
    has_problem = any(w in all_text for w in [
        "masalah", "problem", "kendala", "pain", "isu", "issue",
        "saat ini", "currently", "manual", "tidak bisa", "belum ada",
    ])

    # Section D — Business Objectives: goals
    has_objectives = any(w in all_text for w in [
        "tujuan", "goal", "objective", "target", "ingin", "mau", "want",
        "meningkatkan", "increase", "menurunkan", "reduce",
    ])

    # Section E — Scope: features (in-scope)
    has_features = any(w in all_text for w in [
        "fitur", "feature", "fungsi", "function", "modul", "module",
        "halaman", "page", "dashboard", "login", "register",
    ])

    # Section G — Target Users
    has_users = any(w in all_text for w in [
        "user", "pengguna", "pelanggan", "customer", "target", "audience",
        "admin", "konsumen", "pembeli", "buyer",
    ])

    # Section H — Business Needs: non-trivial requirement detail
    has_requirements = len(all_text) > 300 and any(w in all_text for w in [
        "harus", "must", "perlu", "need", "require", "wajib",
        "sistem", "system", "data", "laporan", "report",
    ])

    # Section K — Risks / Assumptions
    has_risks_or_constraints = any(w in all_text for w in [
        "risiko", "risk", "asumsi", "assumption", "keterbatasan", "constraint",
        "tantangan", "challenge", "hambatan",
    ])

    # Section L — Success Metrics
    has_metrics = any(w in all_text for w in [
        "metrik", "metric", "kpi", "ukur", "measure", "sukses", "success",
        "persentase", "percent", "angka", "number", "target",
    ])

    # Section M — Constraints: budget
    has_budget = any(w in all_text for w in [
        "budget", "biaya", "harga", "anggaran", "rp", "juta", "ribu",
        "million", "cost", "dana",
    ])

    # Section M — Constraints: timeline
    has_timeline = any(w in all_text for w in [
        "deadline", "waktu", "timeline", "kapan", "bulan", "minggu",
        "hari", "day", "week", "month", "selesai", "launch",
    ])

    # Integrations (enriches H and E)
    has_integrations = any(w in all_text for w in [
        "integrasi", "integration", "api", "payment", "pembayaran",
        "whatsapp", "google", "midtrans", "xendit", "notifikasi",
    ])

    checks = [
        has_description,
        has_problem,
        has_objectives,
        has_features,
        has_users,
        has_requirements,
        has_risks_or_constraints,
        has_metrics,
        has_budget,
        has_timeline,
        has_integrations,
    ]

    score = sum(checks) / len(checks) * 100
    return min(100, int(score))


def _score_brd_against_template(brd: dict) -> BrdTemplateScore:
    """Score generated BRD dict against KerjaCUS! BRD template sections (A-N).

    Template sections mapped to BrdDocument fields:
      B  Executive Summary     → executive_summary (length + substance)
      D  Business Objectives   → business_objectives (count + specificity)
      E  Scope                 → scope + out_of_scope
      H  Business Needs/Reqs  → functional_requirements + non_functional_requirements
      K  Risks & Assumptions   → risk_assessment
      L  Success Metrics       → success_metrics
      M  Constraints           → estimated_price_min/max + estimated_timeline_days
    Sections F (Stakeholders), G (Target Users), I (Business Rules),
    J (Expected Benefits), N (Timeline detail) are not in BrdDocument schema —
    marked as gaps with score 0.
    """

    def _score_text(val: object, min_len: int = 100) -> tuple[int, str]:
        if not val:
            return 0, "empty"
        text = str(val)
        if len(text) >= min_len * 2:
            return 100, f"{len(text)} chars"
        if len(text) >= min_len:
            return 70, f"{len(text)} chars (adequate)"
        return 40, f"{len(text)} chars (too brief)"

    def _score_list(val: object, min_items: int = 3, ideal: int = 5) -> tuple[int, str]:
        if not val or not isinstance(val, list):
            return 0, "empty"
        n = len(val)
        if n >= ideal:
            return 100, f"{n} items"
        if n >= min_items:
            return 70, f"{n} items (adequate, aim for {ideal}+)"
        if n >= 1:
            return 40, f"{n} item(s) (too few, need {min_items}+)"
        return 0, "empty list"

    sections: list[BrdSectionScore] = []

    # B — Executive Summary
    s, r = _score_text(brd.get("executive_summary"), min_len=150)
    sections.append(BrdSectionScore(section="B", label="Executive Summary", score=s, reason=r))

    # D — Business Objectives
    s, r = _score_list(brd.get("business_objectives"), min_items=4, ideal=6)
    sections.append(BrdSectionScore(section="D", label="Business Objectives", score=s, reason=r))

    # E — Scope (in-scope)
    s, r = _score_text(brd.get("scope"), min_len=80)
    sections.append(BrdSectionScore(section="E", label="Scope (In-Scope)", score=s, reason=r))

    # E — Scope (out-of-scope)
    s, r = _score_list(brd.get("out_of_scope"), min_items=3, ideal=5)
    sections.append(BrdSectionScore(section="E", label="Scope (Out-of-Scope)", score=s, reason=r))

    # F — Stakeholders (not in schema — always gap)
    sections.append(BrdSectionScore(
        section="F", label="Stakeholders & Roles",
        score=0, reason="Not captured in current BRD schema",
    ))

    # G — Target Users (not in schema — always gap)
    sections.append(BrdSectionScore(
        section="G", label="Target User Segments",
        score=0, reason="Not captured in current BRD schema",
    ))

    # H — Functional Requirements
    s, r = _score_list(brd.get("functional_requirements"), min_items=4, ideal=7)
    sections.append(BrdSectionScore(section="H", label="Functional Requirements", score=s, reason=r))

    # H — Non-Functional Requirements
    s, r = _score_list(brd.get("non_functional_requirements"), min_items=4, ideal=7)
    sections.append(BrdSectionScore(section="H", label="Non-Functional Requirements", score=s, reason=r))

    # I — Business Rules (not in schema — gap)
    sections.append(BrdSectionScore(
        section="I", label="Business Rules",
        score=0, reason="Not captured in current BRD schema",
    ))

    # J — Expected Benefits (not in schema — gap)
    sections.append(BrdSectionScore(
        section="J", label="Expected Benefits",
        score=0, reason="Not captured in current BRD schema",
    ))

    # K — Risks & Assumptions
    s, r = _score_list(brd.get("risk_assessment"), min_items=3, ideal=5)
    sections.append(BrdSectionScore(section="K", label="Risks & Assumptions", score=s, reason=r))

    # L — Success Metrics
    s, r = _score_list(brd.get("success_metrics"), min_items=3, ideal=5)
    sections.append(BrdSectionScore(section="L", label="Success Metrics", score=s, reason=r))

    # M — Budget constraint
    price_min = brd.get("estimated_price_min", 0)
    price_max = brd.get("estimated_price_max", 0)
    if price_min > 0 and price_max > 0:
        bm_score, bm_reason = 100, f"Rp {price_min:,} – Rp {price_max:,}"
    elif price_min > 0 or price_max > 0:
        bm_score, bm_reason = 60, "Partial budget range"
    else:
        bm_score, bm_reason = 0, "No budget estimate"
    sections.append(BrdSectionScore(section="M", label="Budget Estimate", score=bm_score, reason=bm_reason))

    # M — Timeline & team size constraint
    tl = brd.get("estimated_timeline_days", 0)
    ts = brd.get("estimated_team_size", 0)
    if tl > 0 and ts > 0:
        tl_score, tl_reason = 100, f"{tl} days, {ts} person(s)"
    elif tl > 0:
        tl_score, tl_reason = 60, f"{tl} days (team size missing)"
    else:
        tl_score, tl_reason = 0, "No timeline estimate"
    sections.append(BrdSectionScore(section="M", label="Timeline & Team Size", score=tl_score, reason=tl_reason))

    # N — High-level timeline phases (not in schema — gap)
    sections.append(BrdSectionScore(
        section="N", label="High-Level Timeline Phases",
        score=0, reason="Not captured in current BRD schema",
    ))

    total = sum(s.score for s in sections)
    overall = round(total / len(sections)) if sections else 0
    return BrdTemplateScore(overall=overall, sections=sections)


@router.post(
    "/chat",
    response_model=ChatResponse,
    dependencies=[Depends(require_service_auth)],
    responses={502: {"description": "AI gateway unreachable"}},
)
async def chat_completion(request: ChatRequest):
    """AI chatbot for project scoping follow-up. Enriches context via RAG over past BRDs."""
    messages_payload = await _build_chat_messages_with_rag(request)

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{TENSORZERO_URL}/inference",
                json={
                    "function_name": "chatbot",
                    "input": {
                        "messages": messages_payload,
                    },
                },
            )
            response.raise_for_status()
            data = response.json()

        content = data.get("content", [{}])
        text = content[0].get("text", "") if content else ""

        completeness = calculate_completeness(request.messages)

        return ChatResponse(
            message={"role": "assistant", "content": text},
            completeness_score=completeness,
            suggest_generate_brd=completeness >= 80,
        )
    except httpx.HTTPError as e:
        raise HTTPException(status_code=502, detail=f"AI service error: {e}") from e


async def _build_chat_messages_with_rag(request: ChatRequest) -> list[dict]:
    """Construct messages payload, prepending RAG context when available."""
    rag_context_blocks: list[str] = []
    last_user_msg = next(
        (m.content for m in reversed(request.messages) if m.role == "user"),
        "",
    )
    if last_user_msg:
        try:
            from app.services.rag import hybrid_search

            chunks = await hybrid_search(
                query=last_user_msg,
                table="brd_documents",
                content_field="content",
                top_k=4,
            )
            rag_context_blocks = [c["content"] for c in chunks if c.get("content")]
        except Exception as e:
            logger.warning("RAG retrieval failed in /chat: %s", e)

    payload = [m.model_dump() for m in request.messages]
    if rag_context_blocks:
        context_text = "\n\n---\n\n".join(rag_context_blocks)
        payload.insert(
            0,
            {
                "role": "system",
                "content": (
                    "Context from similar past projects (use to ground your "
                    "follow-up questions; do not reveal verbatim):\n"
                    f"{context_text}"
                ),
            },
        )
    return payload


def _sse(data: dict) -> bytes:
    """Encode dict as Server-Sent Event data line."""
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n".encode("utf-8")


async def _stream_chat_tokens(
    request: ChatRequest,
    messages_payload: list[dict],
) -> AsyncIterator[bytes]:
    """Stream TensorZero inference deltas as SSE, finishing with completeness metadata."""
    full_text = ""
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            async with client.stream(
                "POST",
                f"{TENSORZERO_URL}/inference",
                json={
                    "function_name": "chatbot",
                    "input": {"messages": messages_payload},
                    "stream": True,
                },
            ) as response:
                if response.status_code >= 400:
                    body = await response.aread()
                    detail = body.decode("utf-8", errors="ignore")[:500]
                    yield _sse({"type": "error", "message": f"upstream {response.status_code}: {detail}"})
                    return

                async for raw in response.aiter_lines():
                    if not raw:
                        continue
                    line = raw.strip()
                    if not line.startswith("data:"):
                        continue
                    payload = line[5:].strip()
                    if not payload or payload == "[DONE]":
                        continue
                    try:
                        chunk = json.loads(payload)
                    except json.JSONDecodeError:
                        continue
                    delta = _extract_delta_text(chunk)
                    if delta:
                        full_text += delta
                        yield _sse({"type": "token", "delta": delta})
    except httpx.HTTPError as e:
        yield _sse({"type": "error", "message": f"AI gateway error: {e}"})
        return

    completeness = calculate_completeness(request.messages)
    yield _sse(
        {
            "type": "done",
            "full_text": full_text,
            "completeness_score": completeness,
            "suggest_generate_brd": completeness >= 80,
        }
    )


def _extract_delta_text(chunk: dict) -> str:
    """Pull incremental text from a TensorZero stream chunk, tolerating shape drift."""
    content = chunk.get("content")
    if isinstance(content, list) and content:
        first = content[0]
        if isinstance(first, dict):
            for key in ("text", "delta"):
                value = first.get(key)
                if isinstance(value, str) and value:
                    return value
    delta = chunk.get("delta")
    if isinstance(delta, str) and delta:
        return delta
    if isinstance(delta, dict):
        value = delta.get("content") or delta.get("text")
        if isinstance(value, str):
            return value
    return ""


@router.post(
    "/chat/stream",
    dependencies=[Depends(require_service_auth)],
    responses={
        200: {
            "content": {"text/event-stream": {"schema": {}}},
            "description": "Server-Sent Events token stream",
        },
        502: {"description": "AI gateway unreachable"},
    },
)
async def chat_stream(request: ChatRequest):
    """Server-Sent Events stream for chatbot tokens; terminal event carries completeness."""
    messages_payload = await _build_chat_messages_with_rag(request)
    return StreamingResponse(
        _stream_chat_tokens(request, messages_payload),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


def extract_json_from_text(text: str) -> dict:
    """Extract JSON from text that may contain markdown fences."""
    import re

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try extracting from markdown code fence
    match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass

    # Try finding JSON object in text
    brace_start = text.find('{')
    if brace_start >= 0:
        depth = 0
        for i in range(brace_start, len(text)):
            if text[i] == '{':
                depth += 1
            elif text[i] == '}':
                depth -= 1
            if depth == 0:
                try:
                    return json.loads(text[brace_start:i + 1])
                except json.JSONDecodeError:
                    break

    return {}


BRD_SYSTEM_PROMPT = """You are a senior business analyst at KerjaCUS!, a managed marketplace platform for digital projects in Indonesia. Your job is to generate a comprehensive Business Requirement Document (BRD) from the project scoping conversation.

Analyze the conversation history carefully and produce a structured BRD in JSON format with these exact fields:

{
  "executive_summary": "A 2-3 paragraph summary of the project, its goals, target users, and key value proposition.",
  "business_objectives": ["List of 4-6 specific, measurable business objectives"],
  "success_metrics": ["List of 3-5 KPIs to measure project success"],
  "scope": "Detailed paragraph describing what is included in the project scope.",
  "out_of_scope": ["List of 3-5 items explicitly excluded from scope"],
  "functional_requirements": [
    {"title": "Feature Category Name", "content": "Detailed description of the feature and its sub-features"}
  ],
  "non_functional_requirements": ["List of 5-8 NFRs covering performance, security, scalability, accessibility"],
  "estimated_price_min": <integer in IDR>,
  "estimated_price_max": <integer in IDR>,
  "estimated_timeline_days": <integer>,
  "estimated_team_size": <integer>,
  "risk_assessment": ["List of 3-5 key risks with their mitigation strategies, each as a single string in format 'Risk: ... | Mitigation: ...'"]
}

Guidelines:
- Write in English for all technical content.
- Be specific and actionable in requirements -- avoid vague statements.
- Price estimates should be realistic for the Indonesian market (developer rates Rp 15-40 million/month).
- Timeline should account for development, testing, and deployment.
- Team size should match the project complexity and timeline.
- Functional requirements should have 4-8 items covering all major feature areas.
- Always return valid JSON only, no markdown formatting or extra text."""


def _build_brd_messages(
    request: GenerateBrdRequest,
) -> list[dict]:
    """Build the message list for BRD generation from conversation history."""
    messages: list[dict] = [{"role": "system", "content": BRD_SYSTEM_PROMPT}]

    # Add conversation context as a consolidated user message
    conversation_text_parts: list[str] = []
    for msg in request.conversation_history:
        prefix = "Client" if msg.role == "user" else "AI Assistant"
        conversation_text_parts.append(f"{prefix}: {msg.content}")

    context_parts = [
        f"Project Category: {request.project_category}",
    ]
    if request.budget_min is not None:
        context_parts.append(f"Budget Min: Rp {request.budget_min:,}")
    if request.budget_max is not None:
        context_parts.append(f"Budget Max: Rp {request.budget_max:,}")
    if request.timeline_days is not None:
        context_parts.append(f"Requested Timeline: {request.timeline_days} days")

    user_prompt = (
        "Generate a BRD based on the following project scoping conversation and metadata.\n\n"
        f"--- Project Metadata ---\n{chr(10).join(context_parts)}\n\n"
        f"--- Scoping Conversation ---\n{chr(10).join(conversation_text_parts)}\n\n"
        "Return ONLY valid JSON matching the schema described in the system prompt."
    )
    messages.append({"role": "user", "content": user_prompt})
    return messages


def _build_fallback_brd(request: GenerateBrdRequest) -> dict:
    """Build a reasonable BRD from request metadata when LLM fails."""
    # Extract project context from conversation
    conversation_text = " ".join(
        m.content for m in request.conversation_history if m.role == "user"
    )
    summary = conversation_text[:600] if conversation_text else "Digital project"

    budget_min = request.budget_min or 10_000_000
    budget_max = request.budget_max or 50_000_000
    timeline = request.timeline_days or 60
    team_size = max(1, min(5, timeline // 30))

    category_label = request.project_category.replace("_", " ").title()

    return {
        "executive_summary": (
            f"This {category_label} project aims to deliver a digital solution "
            f"based on the client's requirements gathered during scoping. "
            f"The project targets completion within {timeline} days with an estimated "
            f"budget range of Rp {budget_min:,} to Rp {budget_max:,}. "
            f"Key details: {summary[:300]}"
        ),
        "business_objectives": [
            f"Deliver a functional {category_label} within {timeline} days",
            "Meet all core functional requirements defined in scope",
            "Achieve responsive design supporting mobile and desktop",
            "Integrate with required third-party services",
        ],
        "success_metrics": [
            "All defined milestones completed on schedule",
            "Client approval on each milestone deliverable",
            "System passes functional and performance testing",
        ],
        "scope": (
            f"Full development of a {category_label} including frontend, backend, "
            "database design, API development, third-party integrations, "
            "testing, and deployment. Details based on scoping conversation."
        ),
        "out_of_scope": [
            "Native mobile applications (unless specified)",
            "Ongoing maintenance and support post-delivery",
            "Content creation and data migration",
        ],
        "functional_requirements": [
            {
                "title": "Core Application Features",
                "content": "Primary features as discussed during project scoping.",
            },
            {
                "title": "User Management",
                "content": "User registration, authentication, and profile management.",
            },
            {
                "title": "Admin Dashboard",
                "content": "Administrative interface for content and user management.",
            },
        ],
        "non_functional_requirements": [
            "Page load time under 3 seconds on 4G connections",
            "Support for 100+ concurrent users",
            "HTTPS encryption for all data in transit",
            "Responsive design for mobile and desktop",
            "Basic SEO optimization",
        ],
        "estimated_price_min": budget_min,
        "estimated_price_max": budget_max,
        "estimated_timeline_days": timeline,
        "estimated_team_size": team_size,
        "risk_assessment": [
            "Risk: Scope creep from additional requirements | Mitigation: Strict change request process with impact analysis",
            "Risk: Third-party API integration delays | Mitigation: Begin integration early, prepare fallback options",
            "Risk: Timeline pressure affecting quality | Mitigation: Prioritize core features, defer nice-to-haves to Phase 2",
        ],
    }


def _parse_brd_response(text: str, request: GenerateBrdRequest) -> dict:
    """Parse LLM JSON response into BRD dict, falling back on parse errors."""
    parsed = extract_json_from_text(text.strip())
    if not parsed:
        return _build_fallback_brd(request)

    # Normalize functional_requirements: LLM may return {title, content} or {title, description}
    raw_reqs = parsed.get("functional_requirements", [])
    normalized_reqs = []
    for req in raw_reqs:
        if isinstance(req, dict):
            normalized_reqs.append(
                {
                    "title": req.get("title", "Feature"),
                    "content": req.get("content") or req.get("description", ""),
                }
            )
        elif isinstance(req, str):
            normalized_reqs.append({"title": "Requirement", "content": req})

    # Normalize risk_assessment: accept both string list and object list
    raw_risks = parsed.get("risk_assessment", [])
    normalized_risks = []
    for risk in raw_risks:
        if isinstance(risk, str):
            normalized_risks.append(risk)
        elif isinstance(risk, dict):
            r = risk.get("risk", risk.get("title", ""))
            m = risk.get("mitigation", risk.get("strategy", ""))
            normalized_risks.append(f"Risk: {r} | Mitigation: {m}" if m else r)

    fallback = _build_fallback_brd(request)

    return {
        "executive_summary": parsed.get("executive_summary") or fallback["executive_summary"],
        "business_objectives": parsed.get("business_objectives") or fallback["business_objectives"],
        "success_metrics": parsed.get("success_metrics") or fallback["success_metrics"],
        "scope": parsed.get("scope") or fallback["scope"],
        "out_of_scope": parsed.get("out_of_scope") or fallback["out_of_scope"],
        "functional_requirements": normalized_reqs or fallback["functional_requirements"],
        "non_functional_requirements": parsed.get("non_functional_requirements") or fallback["non_functional_requirements"],
        "estimated_price_min": parsed.get("estimated_price_min") or fallback["estimated_price_min"],
        "estimated_price_max": parsed.get("estimated_price_max") or fallback["estimated_price_max"],
        "estimated_timeline_days": parsed.get("estimated_timeline_days") or fallback["estimated_timeline_days"],
        "estimated_team_size": parsed.get("estimated_team_size") or fallback["estimated_team_size"],
        "risk_assessment": normalized_risks or fallback["risk_assessment"],
    }


@router.post(
    "/generate-brd",
    response_model=GenerateBrdResponse,
    dependencies=[Depends(require_service_auth)],
    responses={502: {"description": "AI gateway unreachable"}},
)
async def generate_brd(request: GenerateBrdRequest):
    """Generate BRD from conversation history via TensorZero LLM gateway."""
    messages = _build_brd_messages(request)
    model_used = "gemini-pro"
    tokens_used = 0

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                f"{TENSORZERO_URL}/inference",
                json={
                    "function_name": "brd_generation",
                    "input": {
                        "messages": messages,
                    },
                },
            )
            response.raise_for_status()
            data = response.json()

        content_blocks = data.get("content", [{}])
        text = content_blocks[0].get("text", "") if content_blocks else ""

        usage = data.get("usage", {})
        tokens_used = usage.get("input_tokens", 0) + usage.get("output_tokens", 0)
        model_used = data.get("model", model_used)

        brd = _parse_brd_response(text, request)

    except (httpx.HTTPError, KeyError, IndexError):
        # TensorZero unavailable or returned unexpected shape -- use fallback
        brd = _build_fallback_brd(request)

    template_score = _score_brd_against_template(brd)

    await publish_event(
        "ai.brd.generated",
        {
            "projectId": request.project_id,
            "tokensUsed": tokens_used,
            "model": model_used,
            "templateScore": template_score.overall,
        },
    )

    return GenerateBrdResponse(
        brd=brd,
        tokens_used=tokens_used,
        model=model_used,
        template_score=template_score,
    )


PRD_SYSTEM_PROMPT = """You are a senior technical architect at KerjaCUS!, a managed marketplace platform for digital projects in Indonesia. Your job is to generate a comprehensive Product Requirement Document (PRD) from the BRD and project context.

Analyze the BRD content and conversation history carefully and produce a structured PRD in JSON format with these exact fields:

{
  "tech_stack": ["List of recommended technologies, e.g. React, Node.js, PostgreSQL"],
  "architecture": "Detailed paragraph describing the system architecture (microservices, monolith, serverless, etc.) with justification.",
  "api_design": "Description of the API design approach (REST, GraphQL, etc.) with key endpoints listed.",
  "database_schema": "Description of the database design with key tables, relationships, and indexing strategy.",
  "team_composition": {
    "team_size": <integer>,
    "work_packages": [
      {
        "title": "Work Package Name (e.g. Frontend Development)",
        "description": "Detailed description of the work package scope",
        "required_skills": ["skill1", "skill2"],
        "estimated_hours": <float>,
        "amount": <integer in IDR>
      }
    ]
  },
  "work_packages": [<same as team_composition.work_packages>],
  "sprint_plan": [
    {
      "sprint_number": <integer>,
      "title": "Sprint Title",
      "tasks": ["Task 1 description", "Task 2 description"],
      "duration_days": <integer, typically 14>
    }
  ],
  "dependencies": [
    {
      "from_package": "Work Package Title that must finish first",
      "to_package": "Work Package Title that depends on it",
      "type": "finish_to_start"
    }
  ],
  "estimated_price_min": <integer in IDR>,
  "estimated_price_max": <integer in IDR>,
  "estimated_timeline_days": <integer>,
  "estimated_team_size": <integer>
}

Guidelines:
- Write in English for all technical content.
- Tech stack should be specific (versions if relevant) and justified for the project type.
- Team size calculation: total_estimated_hours / (timeline_days * 6 working_hours_per_day), minimum 1, maximum 8.
- Work packages should be decomposed by role/skill area (Frontend, Backend, UI/UX, etc.).
- Sprint plan should have 2-week sprints covering the full timeline.
- Dependencies should form a valid DAG (no cycles).
- Pricing should be realistic for the Indonesian market.
- Always return valid JSON only, no markdown formatting or extra text."""


def _build_prd_messages(request: GeneratePrdRequest) -> list[dict]:
    """Build the message list for PRD generation from BRD and conversation."""
    messages: list[dict] = [{"role": "system", "content": PRD_SYSTEM_PROMPT}]

    context_parts = [
        f"Project Category: {request.project_category}",
    ]
    if request.budget_min is not None:
        context_parts.append(f"Budget Min: Rp {request.budget_min:,}")
    if request.budget_max is not None:
        context_parts.append(f"Budget Max: Rp {request.budget_max:,}")
    if request.timeline_days is not None:
        context_parts.append(f"Requested Timeline: {request.timeline_days} days")

    conversation_text_parts: list[str] = []
    for msg in request.conversation_history:
        prefix = "Client" if msg.role == "user" else "AI Assistant"
        conversation_text_parts.append(f"{prefix}: {msg.content}")

    brd_json = json.dumps(request.brd_content, indent=2, default=str)

    user_prompt = (
        "Generate a PRD based on the following BRD document and project metadata.\n\n"
        f"--- Project Metadata ---\n{chr(10).join(context_parts)}\n\n"
        f"--- BRD Document ---\n{brd_json}\n\n"
    )
    if conversation_text_parts:
        user_prompt += f"--- Scoping Conversation ---\n{chr(10).join(conversation_text_parts)}\n\n"
    user_prompt += "Return ONLY valid JSON matching the schema described in the system prompt."

    messages.append({"role": "user", "content": user_prompt})
    return messages


def _build_fallback_prd(request: GeneratePrdRequest) -> dict:
    """Build a reasonable PRD from BRD data when LLM fails."""
    brd = request.brd_content
    budget_min = request.budget_min or brd.get("estimated_price_min", 10_000_000)
    budget_max = request.budget_max or brd.get("estimated_price_max", 50_000_000)
    timeline = request.timeline_days or brd.get("estimated_timeline_days", 60)
    team_size = brd.get("estimated_team_size", max(1, min(5, timeline // 30)))

    category_label = request.project_category.replace("_", " ").title()

    # Default work packages based on category
    work_packages = [
        {
            "title": "Backend API Development",
            "description": f"Server-side logic, API endpoints, database integration for {category_label}",
            "required_skills": ["Node.js", "PostgreSQL", "REST API"],
            "estimated_hours": float(timeline * 4),
            "amount": int(budget_min * 0.35),
        },
        {
            "title": "Frontend Development",
            "description": f"User interface implementation for {category_label}",
            "required_skills": ["React", "TypeScript", "Tailwind CSS"],
            "estimated_hours": float(timeline * 4),
            "amount": int(budget_min * 0.35),
        },
        {
            "title": "UI/UX Design",
            "description": "Wireframes, mockups, design system, and prototypes",
            "required_skills": ["Figma", "UI Design", "UX Research"],
            "estimated_hours": float(timeline * 2),
            "amount": int(budget_min * 0.2),
        },
    ]

    sprints = []
    num_sprints = max(1, timeline // 14)
    for i in range(num_sprints):
        sprints.append(
            {
                "sprint_number": i + 1,
                "title": f"Sprint {i + 1}",
                "tasks": [f"Development tasks for sprint {i + 1}"],
                "duration_days": 14,
            }
        )

    return {
        "tech_stack": ["React", "TypeScript", "Node.js", "PostgreSQL", "Tailwind CSS", "Docker"],
        "architecture": (
            f"Modular monolith architecture for {category_label} with clear service boundaries. "
            "REST API backend with PostgreSQL database, React frontend with server-side rendering support."
        ),
        "api_design": (
            "RESTful API design with versioned endpoints (/api/v1/*). "
            "JSON request/response format, JWT authentication, pagination for list endpoints."
        ),
        "database_schema": (
            "Normalized PostgreSQL schema with UUID primary keys, timestamptz for all timestamps. "
            "Key tables based on BRD functional requirements with proper indexing and foreign key constraints."
        ),
        "team_composition": {
            "team_size": team_size,
            "work_packages": work_packages,
        },
        "work_packages": work_packages,
        "sprint_plan": sprints,
        "dependencies": [
            {
                "from_package": "UI/UX Design",
                "to_package": "Frontend Development",
                "type": "finish_to_start",
            },
            {
                "from_package": "Backend API Development",
                "to_package": "Frontend Development",
                "type": "start_to_start",
            },
        ],
        "estimated_price_min": budget_min,
        "estimated_price_max": budget_max,
        "estimated_timeline_days": timeline,
        "estimated_team_size": team_size,
    }


def _parse_prd_response(text: str, request: GeneratePrdRequest) -> dict:
    """Parse LLM JSON response into PRD dict, falling back on parse errors."""
    parsed = extract_json_from_text(text.strip())
    if not parsed:
        return _build_fallback_prd(request)

    fallback = _build_fallback_prd(request)

    # Normalize work_packages
    raw_wps = parsed.get("work_packages", [])
    normalized_wps = []
    for wp in raw_wps:
        if isinstance(wp, dict):
            normalized_wps.append(
                {
                    "title": wp.get("title", "Work Package"),
                    "description": wp.get("description", ""),
                    "required_skills": wp.get("required_skills", []),
                    "estimated_hours": float(wp.get("estimated_hours", 0)),
                    "amount": int(wp.get("amount", 0)),
                }
            )

    # Normalize sprint_plan
    raw_sprints = parsed.get("sprint_plan", [])
    normalized_sprints = []
    for sp in raw_sprints:
        if isinstance(sp, dict):
            normalized_sprints.append(
                {
                    "sprint_number": int(sp.get("sprint_number", len(normalized_sprints) + 1)),
                    "title": sp.get("title", f"Sprint {len(normalized_sprints) + 1}"),
                    "tasks": sp.get("tasks", []),
                    "duration_days": int(sp.get("duration_days", 14)),
                }
            )

    # Normalize dependencies
    raw_deps = parsed.get("dependencies", [])
    normalized_deps = []
    for dep in raw_deps:
        if isinstance(dep, dict):
            normalized_deps.append(
                {
                    "from_package": dep.get("from_package", ""),
                    "to_package": dep.get("to_package", ""),
                    "type": dep.get("type", "finish_to_start"),
                }
            )

    # Normalize team_composition
    raw_tc = parsed.get("team_composition", {})
    team_composition = {
        "team_size": raw_tc.get("team_size", parsed.get("estimated_team_size", fallback["estimated_team_size"])),
        "work_packages": normalized_wps or fallback["work_packages"],
    }

    return {
        "tech_stack": parsed.get("tech_stack") or fallback["tech_stack"],
        "architecture": parsed.get("architecture") or fallback["architecture"],
        "api_design": parsed.get("api_design") or fallback["api_design"],
        "database_schema": parsed.get("database_schema") or fallback["database_schema"],
        "team_composition": team_composition,
        "work_packages": normalized_wps or fallback["work_packages"],
        "sprint_plan": normalized_sprints or fallback["sprint_plan"],
        "dependencies": normalized_deps or fallback["dependencies"],
        "estimated_price_min": parsed.get("estimated_price_min") or fallback["estimated_price_min"],
        "estimated_price_max": parsed.get("estimated_price_max") or fallback["estimated_price_max"],
        "estimated_timeline_days": parsed.get("estimated_timeline_days") or fallback["estimated_timeline_days"],
        "estimated_team_size": parsed.get("estimated_team_size") or fallback["estimated_team_size"],
    }


@router.post(
    "/generate-prd",
    response_model=GeneratePrdResponse,
    dependencies=[Depends(require_service_auth)],
    responses={502: {"description": "AI gateway unreachable"}},
)
async def generate_prd(request: GeneratePrdRequest):
    """Generate PRD from BRD content and conversation history via TensorZero LLM gateway."""
    messages = _build_prd_messages(request)
    model_used = "gemini-pro"
    tokens_used = 0

    try:
        async with httpx.AsyncClient(timeout=90.0) as client:
            response = await client.post(
                f"{TENSORZERO_URL}/inference",
                json={
                    "function_name": "prd_generation",
                    "input": {
                        "messages": messages,
                    },
                },
            )
            response.raise_for_status()
            data = response.json()

        content_blocks = data.get("content", [{}])
        text = content_blocks[0].get("text", "") if content_blocks else ""

        usage = data.get("usage", {})
        tokens_used = usage.get("input_tokens", 0) + usage.get("output_tokens", 0)
        model_used = data.get("model", model_used)

        prd = _parse_prd_response(text, request)

    except (httpx.HTTPError, KeyError, IndexError):
        prd = _build_fallback_prd(request)

    await publish_event(
        "ai.prd.generated",
        {
            "projectId": request.project_id,
            "tokensUsed": tokens_used,
            "model": model_used,
        },
    )

    return GeneratePrdResponse(
        prd=prd,
        tokens_used=tokens_used,
        model=model_used,
    )


@router.post(
    "/parse-cv",
    response_model=CvParseResponse,
)
async def parse_cv(request: CvParseRequest):
    """Parse CV using document text extraction + LLM structured extraction via Instructor."""
    import asyncio
    import tempfile
    from pathlib import Path

    import instructor
    from openai import OpenAI
    from pydantic import BaseModel, Field

    class ExtractedCV(BaseModel):
        name: str = Field(default="", description="Full name")
        email: str = Field(default="", description="Email address")
        phone: str = Field(default="", description="Phone number")
        summary: str = Field(default="", description="Professional summary or objective statement")
        skills: list[str] = Field(
            default_factory=list,
            description=(
                "ALL technical skills extracted from EVERY section: "
                "certifications tech tags, project tech stacks, work experience descriptions, "
                "education coursework, and any explicit skills section. "
                "Include frameworks, languages, tools, platforms, algorithms, and ML model types."
            ),
        )
        education: list[dict] = Field(
            default_factory=list,
            description="Education history. Each item: {university, major, year, gpa}",
        )
        experience: list[dict] = Field(
            default_factory=list,
            description="Work experience. Each item: {company, position, start, end, description}",
        )
        organizational_experience: list[dict] = Field(
            default_factory=list,
            description="Organizational/volunteer experience. Each item: {organization, role, start, end, description}",
        )
        projects: list[dict] = Field(
            default_factory=list,
            description="Personal/academic projects. Each item: {title, tech_stack, description, url}",
        )
        certifications: list[dict] = Field(
            default_factory=list,
            description="Certifications. Each item: {name, issuer, year}",
        )
        portfolio_urls: list[str] = Field(
            default_factory=list,
            description="Portfolio/professional URLs (GitHub, LinkedIn, Dribbble, Behance, etc.)",
        )
        years_of_experience: int | None = Field(
            default=None,
            description="Total years of professional work experience (integer, exclude internships if under 1 year)",
        )

    raw_file_url = request.file_url or ""
    if raw_file_url.startswith(("http://", "https://")):
        file_url = raw_file_url
    else:
        s3_url = os.getenv("S3_ENDPOINT", "http://localhost:9000")
        bucket = os.getenv("S3_BUCKET", "kerjacus-uploads")
        file_url = f"{s3_url.rstrip('/')}/{bucket}/{raw_file_url.lstrip('/')}"

    file_bytes = None
    for attempt in range(3):
        try:
            async with httpx.AsyncClient(timeout=15.0) as dl:
                res = await dl.get(file_url)
                if res.status_code == 200:
                    file_bytes = res.content
                    break
                logger.warning(
                    "CV download attempt %d failed: status=%d url=%s",
                    attempt + 1, res.status_code, file_url,
                )
        except Exception as e:
            logger.warning("CV download attempt %d errored: %s", attempt + 1, e)
        await asyncio.sleep(1)

    # Step 2: Extract text based on file type
    cv_text = ""

    if file_bytes:
        ext = (request.file_type or "pdf").lower()
        try:
            if ext == "pdf":
                try:
                    import pypdfium2 as pdfium
                    pdf = pdfium.PdfDocument(file_bytes)
                    pages = []
                    for page in pdf:
                        textpage = page.get_textpage()
                        pages.append(textpage.get_text_bounded())
                        textpage.close()
                        page.close()
                    pdf.close()
                    cv_text = "\n".join(pages)
                except Exception:
                    cv_text = file_bytes.decode("utf-8", errors="ignore")
            elif ext in ("docx", "doc"):
                tmp_path = None
                try:
                    import docx
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(file_bytes)
                        tmp_path = tmp.name
                    doc = docx.Document(tmp_path)
                    cv_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except Exception:
                    cv_text = file_bytes.decode("utf-8", errors="ignore")
                finally:
                    if tmp_path:
                        Path(tmp_path).unlink(missing_ok=True)
            else:
                cv_text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            cv_text = file_bytes.decode("utf-8", errors="ignore")

    if not cv_text or len(cv_text.strip()) < 50:
        return CvParseResponse(
            talent_id=request.talent_id,
            parsed_data=CvParsedData(),
            confidence_score=0.0,
            raw_text="",
        )

    # Step 3: Use Instructor for structured extraction
    tensorzero_url = os.getenv("TENSORZERO_API_URL", "http://localhost:3333")
    llm_api_key = os.getenv("LLM_API_KEY", "")

    try:
        client = instructor.from_openai(
            OpenAI(api_key=llm_api_key, base_url=f"{tensorzero_url}/openai/v1"),
        )

        extracted = client.chat.completions.create(
            model="cv_extraction",
            response_model=ExtractedCV,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an expert CV parser. Extract ALL structured information from this CV/resume text.\n\n"
                        "CRITICAL — skills extraction rules:\n"
                        "1. Scan EVERY section: certifications tech tags, project tech stacks (after '|' or in parentheses), "
                        "work experience bullet points, education coursework, and any explicit skills section.\n"
                        "2. Include: programming languages, frameworks, libraries, tools, platforms, cloud services, "
                        "databases, ML algorithms (XGBoost, CatBoost, LightGBM, KNN, GNB, CNN, LSTM, etc.), "
                        "MLOps tools (MLflow, Kubeflow, KServe, Feast), data tools (Tableau, R, Streamlit, Gradio), "
                        "AI frameworks (LangChain, FAISS, Transformers, Hugging Face, LLM).\n"
                        "3. Do NOT invent skills not mentioned in the text.\n"
                        "4. Deduplicate — return each skill once.\n\n"
                        "For organizational_experience: extract leadership roles in student orgs, volunteer work, committees.\n"
                        "For projects: extract from 'Projects' section — each with title, tech stack list, short description, and any URL.\n"
                        "For certifications: include the tech tags listed after the cert name (e.g. 'IBM AI Engineering | Python, PyTorch, TensorFlow').\n"
                        "For Indonesian CVs, handle both Indonesian and English content."
                    ),
                },
                {"role": "user", "content": cv_text[:12000]},
            ],
            max_retries=2,
        )

        parsed_data = CvParsedData(
            name=extracted.name,
            email=extracted.email,
            phone=extracted.phone,
            summary=extracted.summary or None,
            skills=extracted.skills,
            education=extracted.education,
            experience=extracted.experience,
            organizational_experience=extracted.organizational_experience,
            projects=extracted.projects,
            certifications=extracted.certifications,
            portfolio_urls=extracted.portfolio_urls,
            years_of_experience=extracted.years_of_experience,
        )

        # Confidence based on field completeness
        filled_fields = sum(1 for v in [
            extracted.name, extracted.email, extracted.phone,
            extracted.skills, extracted.education, extracted.experience,
        ] if v)
        confidence = min(0.95, 0.3 + (filled_fields / 6) * 0.65)

    except Exception:
        # Fallback to regex-based parsing
        from app.services.cv_parser import parse_cv_text

        result = parse_cv_text(cv_text)
        parsed_data = CvParsedData(
            name=result.name,
            email=result.email,
            phone=result.phone,
            skills=result.skills,
            education=result.education,
            experience=result.experience,
            projects=result.projects,
            portfolio_urls=result.portfolio_urls,
        )
        confidence = min(0.7, 0.3 + len(result.skills) * 0.04)

    await publish_event(
        "ai.cv.parsed",
        {
            "talentId": request.talent_id,
            "confidenceScore": float(confidence),
            "skillCount": len(parsed_data.skills or []),
        },
    )

    return CvParseResponse(
        talent_id=request.talent_id,
        parsed_data=parsed_data,
        confidence_score=confidence,
        raw_text=cv_text[:2000],
    )


SPEC_PARSE_SYSTEM_PROMPT = """You are a project specification analyzer for KerjaCUS!, a managed marketplace for digital projects in Indonesia.
Extract key project information from the uploaded specification document.

Return a JSON object with exactly these fields:
{
  "summary": "A 2-3 sentence summary of the project",
  "features": ["list of key features or requirements mentioned"],
  "target_users": "description of intended users or audience",
  "integrations": ["list of third-party systems or APIs mentioned"],
  "tech_requirements": "any technical requirements, constraints, or preferences mentioned",
  "budget_hints": "any budget, cost, or pricing information mentioned (empty string if none)",
  "timeline_hints": "any timeline, deadline, or schedule information mentioned (empty string if none)",
  "completeness": <integer 0-100, how complete this spec is for generating a Business Requirements Document>
}

The completeness score should reflect how much information is available for generating a BRD:
- 90-100: Very detailed spec with features, users, tech, budget, timeline
- 70-89: Good spec with most key areas covered
- 50-69: Partial spec, missing several important areas
- 30-49: Brief overview, needs significant follow-up
- 0-29: Very sparse, barely any useful information

Return ONLY valid JSON, no markdown or extra text."""


@router.post(
    "/parse-spec",
    response_model=ParseSpecResponse,
    dependencies=[Depends(require_service_auth)],
)
async def parse_spec(request: ParseSpecRequest):
    """Parse an uploaded specification document and extract project information."""
    file_url = request.file_url
    file_type = request.file_type
    notes = request.notes

    # Download file
    file_bytes = None
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            res = await client.get(file_url)
            if res.status_code == 200:
                file_bytes = res.content
    except Exception:
        pass

    if not file_bytes and not file_url.startswith(("http://", "https://")):
        s3_url = os.getenv("S3_ENDPOINT", "http://localhost:9000")
        bucket = os.getenv("S3_BUCKET", "kerjacus-uploads")
        s3_file_url = f"{s3_url.rstrip('/')}/{bucket}/{file_url.lstrip('/')}"
        try:
            async with httpx.AsyncClient(timeout=15.0) as client:
                res = await client.get(s3_file_url)
                if res.status_code == 200:
                    file_bytes = res.content
                else:
                    logger.warning("Spec S3 download failed: status=%d", res.status_code)
        except Exception as e:
            logger.warning("Spec S3 download errored: %s", e)

    if not file_bytes:
        return ParseSpecResponse(
            data=ParseSpecData(
                summary="Failed to download specification file.",
                completeness=0,
            ),
        )

    # Parse document text
    from app.services.cv_parser import extract_text

    raw_text = extract_text(file_bytes, file_type)

    if not raw_text or len(raw_text.strip()) < 50:
        return ParseSpecResponse(
            data=ParseSpecData(
                summary="Document too short to extract meaningful information.",
                completeness=10,
            ),
        )

    # Extract project information using LLM
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            ai_res = await client.post(
                f"{TENSORZERO_URL}/inference",
                json={
                    "function_name": "chatbot",
                    "input": {
                        "messages": [
                            {"role": "system", "content": SPEC_PARSE_SYSTEM_PROMPT},
                            {
                                "role": "user",
                                "content": f"Parse this specification document:\n\n{raw_text[:8000]}\n\nAdditional notes from the client: {notes}",
                            },
                        ],
                    },
                },
            )
            if ai_res.status_code == 200:
                ai_data = ai_res.json()
                content = ai_data.get("content", [{}])
                text = content[0].get("text", "{}") if content else "{}"
                parsed = extract_json_from_text(text)

                return ParseSpecResponse(
                    data=ParseSpecData(
                        summary=parsed.get("summary", raw_text[:500]),
                        features=parsed.get("features", []),
                        target_users=parsed.get("target_users", ""),
                        integrations=parsed.get("integrations", []),
                        tech_requirements=parsed.get("tech_requirements", ""),
                        budget_hints=parsed.get("budget_hints", ""),
                        timeline_hints=parsed.get("timeline_hints", ""),
                        completeness=min(100, max(0, int(parsed.get("completeness", 50)))),
                    ),
                )
    except Exception:
        pass

    # Fallback: return raw text summary
    return ParseSpecResponse(
        data=ParseSpecData(
            summary=raw_text[:500],
            completeness=40,
        ),
    )


@router.post(
    "/match-talents",
    response_model=MatchingResponse,
    dependencies=[Depends(require_service_auth)],
)
async def match_talents(request: MatchingRequest):
    """Match talents to a project: delegate scoring to project-service rule-based recommender."""
    headers = {}
    secret = _service_auth_secret()
    if secret:
        headers["X-Service-Auth"] = secret

    project_recommendations: list[dict] = []
    exploration_count = 0
    exploitation_count = 0
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            res = await client.post(
                f"{PROJECT_SERVICE_URL}/api/v1/matching/recommend",
                json={
                    "requiredSkills": request.required_skills,
                    "limit": 10,
                },
                headers=headers,
            )
            res.raise_for_status()
            payload = res.json()
            data = payload.get("data", {}) if isinstance(payload, dict) else {}
            project_recommendations = data.get("recommendations", [])
            exploration_count = data.get("explorationCount", 0)
            exploitation_count = data.get("exploitationCount", 0)
    except httpx.HTTPError as e:
        logger.warning("project-service matching unavailable: %s", e)
        return MatchingResponse(
            project_id=request.project_id,
            recommendations=[],
            exploration_count=0,
            exploitation_count=0,
        )

    recommendations = []
    for r in project_recommendations:
        recommendations.append(
            {
                "talent_id": r.get("talentId", ""),
                "score": float(r.get("score", 0)),
                "skill_match": float(r.get("skillMatch", 0)),
                "pemerataan_score": float(r.get("pemerataanScore", 0)),
                "track_record": float(r.get("trackRecord", 0)),
                "rating": float(r.get("rating", 0)),
                "is_exploration": bool(r.get("isExploration", False)),
            }
        )

    await publish_event(
        "ai.matching.completed",
        {
            "projectId": request.project_id,
            "recommendationCount": len(recommendations),
            "explorationCount": exploration_count,
            "exploitationCount": exploitation_count,
        },
    )

    return MatchingResponse(
        project_id=request.project_id,
        recommendations=recommendations,
        exploration_count=exploration_count,
        exploitation_count=exploitation_count,
    )


class EmbedDocumentRequest(BaseModel):
    documentId: str
    documentType: Literal["brd", "prd"]
    content: str | dict | list


@router.post(
    "/embed-document",
    dependencies=[Depends(require_service_auth)],
    responses={
        500: {"description": "Embedding or persistence error"},
        503: {"description": "Embedding service unavailable"},
    },
)
async def embed_document(request: EmbedDocumentRequest):
    """Compute Gemini embedding for a BRD/PRD and persist it to the document row.

    Internal endpoint -- expected to be called from project-service after approval.
    """
    document_id = request.documentId
    document_type = request.documentType
    content = request.content

    if isinstance(content, dict | list):
        text_input = json.dumps(content, default=str)[:8000]
    else:
        text_input = str(content)[:8000]

    table = "brd_documents" if document_type == "brd" else "prd_documents"

    try:
        from app.services.embedding import embed_text
        from app.services.rag import write_embedding

        embedding = await embed_text(text_input)
        ok = await write_embedding(table=table, row_id=document_id, embedding=embedding)
        if not ok:
            raise HTTPException(status_code=500, detail="Failed to persist embedding")
    except HTTPException:
        raise
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except Exception as e:
        logger.exception("embed-document failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Embedding error: {e}") from e

    return {"success": True, "documentId": document_id, "dimensions": len(embedding)}
